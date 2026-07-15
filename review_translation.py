#!/usr/bin/env python3
"""
翻訳品質チェックツール(誤訳・翻訳漏れの検出)
=====================================================================
原文と翻訳文(PDF / Word / テキスト)を読み込んで比較し、
誤訳・翻訳漏れ・数値の不一致などを検出してレポートを出力します。
原文・翻訳文がスキャン画像PDF(テキスト層なし)の場合は、
Claudeのvision機能で自動的に転写(OCR)してから比較します。

2段階でチェックします:
  1. 機械チェック(Python): 条番号の欠落、数値(金額・日付等)の不一致、
     判読不能マーカーの残存を決定論的に検出
  2. AIレビュー(Claude Opus): 誤訳・翻訳漏れ・不要な追加・
     用語の不統一・法的ニュアンスのずれを詳細にレビュー

日本語⇔英語のどちらの方向でも使用できます(自動判定)。

使い方:
    1. 下の「★★★ 設定 ★★★」ブロックを編集
    2. 実行: python review_translation.py
"""

import base64
import io
import json
import os
import re
import sys
import time
from collections import Counter
from pathlib import Path

import anthropic
from dotenv import load_dotenv

load_dotenv()

# ===========================================================================
# ★★★ 設定(ここを編集してください)★★★
# ===========================================================================
SOURCE_FILE     = "samples/sample_contract_ja_scan.pdf"   # 原文ファイル
TRANSLATED_FILE = "output/sample_contract_ja_scan_EN.pdf" # 翻訳後ファイル
OUTPUT_DIR      = "output"                           # レポートの出力先
# ===========================================================================

# ---------------------------------------------------------------------------
# 内部設定
# ---------------------------------------------------------------------------
DEFAULT_MODEL = "claude-opus-4-8"   # 高精度レビュー用(Opus)
TRANSCRIBE_MODEL = "claude-sonnet-5"  # スキャンPDFの転写(OCR)用(コスト重視)
TRANSCRIBE_PAGES_PER_BATCH = 10     # 転写時に1回で処理するページ数
MAX_TOKENS_PER_CALL = 16384
MAX_RETRIES = 3
MAX_TOTAL_CHARS = 150_000  # 原文+翻訳文の合計がこれを超える場合は警告

SYSTEM_PROMPT = """You are a senior bilingual legal reviewer specializing in \
Japanese-English contract translation quality assurance. You are meticulous \
and never miss omissions or mistranslations."""

REVIEW_PROMPT = """Below are a source contract and its translation \
(Japanese-English or English-Japanese; detect the direction yourself).

Review the translation exhaustively, clause by clause, and identify:
1. 誤訳 (mistranslation): meaning differs from the source
2. 翻訳漏れ (omission): source content missing from the translation \
(check EVERY clause, item, date, name, address, and signature block)
3. 不要な追加 (addition): content not present in the source
4. 数値・日付の不一致 (number/date mismatch): amounts, dates, article numbers, periods
5. 用語の不統一 (terminology inconsistency): e.g. 甲/乙 vs Party A/B mapping, \
defined terms used inconsistently
6. 法的ニュアンスのずれ (legal nuance): shall/may, 義務/努力義務 の混同など

Respond ONLY with a JSON object in this exact format (no code fences, \
no commentary). Write all descriptions and suggestions in Japanese:

{
  "overall_assessment": "全体評価を2〜3文で",
  "quality_score": <0-100の整数>,
  "issues": [
    {
      "severity": "重大" | "中" | "軽微",
      "type": "誤訳" | "翻訳漏れ" | "追加" | "数値不一致" | "用語不統一" | "法的ニュアンス",
      "location": "該当箇所(例: 第5条 / Article 5)",
      "source_excerpt": "原文の該当部分(短く)",
      "translation_excerpt": "翻訳の該当部分(短く。漏れの場合は空文字)",
      "description": "問題の説明",
      "suggestion": "修正案"
    }
  ]
}

If there are no issues, return an empty "issues" array.

=== SOURCE DOCUMENT ===
{SOURCE}

=== TRANSLATION ===
{TRANSLATION}"""


# ---------------------------------------------------------------------------
# 1. ファイル読み込み(PDF / docx / txt)
# ---------------------------------------------------------------------------
def read_source(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            text = "\n\n".join((p.extract_text() or "") for p in pdf.pages).strip()
        return text  # スキャンPDFの場合は空文字を返す(呼び出し側で転写にフォールバック)
    if suffix == ".docx":
        from docx import Document
        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells))
        return "\n".join(parts).strip()
    if suffix in (".txt", ".text", ".md"):
        for enc in ("utf-8", "utf-8-sig", "cp932", "shift_jis", "euc-jp",
                    "cp1252", "latin-1"):
            try:
                return path.read_text(encoding=enc).strip()
            except (UnicodeDecodeError, UnicodeError):
                continue
        raise ValueError(f"{path.name}: 文字コードを判別できませんでした。")
    raise ValueError(f"未対応のファイル形式です: {suffix}")


def is_japanese(text: str) -> bool:
    jp = len(re.findall(r"[\u3041-\u30FF\u4E00-\u9FFF]", text))
    return jp > len(text) * 0.1


# ---------------------------------------------------------------------------
# 1.5 スキャンPDFの転写(OCR)フォールバック
# ---------------------------------------------------------------------------
TRANSCRIBE_PROMPT = (
    "The attached PDF is a scanned document. Transcribe ALL text exactly as "
    "written, preserving the original language, structure, line breaks, "
    "clause numbering, dates, amounts, names, and addresses. Do NOT "
    "translate, summarize, or correct anything. If a part is illegible, "
    "write [illegible] at that position. Output ONLY the transcription, "
    "with no preamble or commentary."
)


def get_client() -> anthropic.Anthropic:
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        sys.exit(
            "エラー: ANTHROPIC_API_KEY が見つかりません。\n"
            "  .env ファイルに ANTHROPIC_API_KEY=sk-ant-... を記述してください。"
        )
    return anthropic.Anthropic(api_key=api_key)


def _pdf_batches(path: Path, pages_per_batch: int) -> list[bytes]:
    from pypdf import PdfReader, PdfWriter
    reader = PdfReader(str(path))
    batches = []
    for start in range(0, len(reader.pages), pages_per_batch):
        writer = PdfWriter()
        for page in reader.pages[start: start + pages_per_batch]:
            writer.add_page(page)
        buf = io.BytesIO()
        writer.write(buf)
        batches.append(buf.getvalue())
    return batches


def transcribe_scanned_pdf(path: Path, client: anthropic.Anthropic) -> str:
    """テキスト層のないPDFをvision機能で転写する(翻訳はしない)。"""
    batches = _pdf_batches(path, TRANSCRIBE_PAGES_PER_BATCH)
    print(f"  スキャンPDFを転写中({len(batches)} バッチ、"
          f"モデル: {TRANSCRIBE_MODEL})...")
    parts = []
    for i, pdf_bytes in enumerate(batches, 1):
        pdf_b64 = base64.standard_b64encode(pdf_bytes).decode("utf-8")
        last_error = None
        for attempt in range(1, MAX_RETRIES + 1):
            try:
                with client.messages.stream(
                    model=TRANSCRIBE_MODEL,
                    max_tokens=MAX_TOKENS_PER_CALL,
                    messages=[{
                        "role": "user",
                        "content": [
                            {"type": "document",
                             "source": {"type": "base64",
                                        "media_type": "application/pdf",
                                        "data": pdf_b64}},
                            {"type": "text", "text": TRANSCRIBE_PROMPT},
                        ],
                    }],
                ) as stream:
                    parts.append(stream.get_final_text().strip())
                break
            except anthropic.AuthenticationError:
                sys.exit(
                    "エラー: APIキーが無効です(401)。"
                    ".env の ANTHROPIC_API_KEY を確認してください。"
                )
            except anthropic.APIError as e:
                last_error = e
                wait = 2 ** attempt
                print(f"  [警告] API エラー(試行 {attempt}/{MAX_RETRIES}): {e}")
                time.sleep(wait)
        else:
            raise RuntimeError(f"転写に失敗しました: {last_error}")
        print(f"  転写バッチ {i}/{len(batches)} 完了")
    return "\n".join(parts)


def load_document(path: Path, client_holder: dict) -> str:
    """ファイルを読み込む。テキスト層のないPDFは自動的に転写する。"""
    text = read_source(path)
    if not text and path.suffix.lower() == ".pdf":
        print(f"  {path.name}: テキスト層なし → visionで転写(OCR)します")
        if client_holder.get("client") is None:
            client_holder["client"] = get_client()
        text = transcribe_scanned_pdf(path, client_holder["client"])
        if not text:
            sys.exit(f"エラー: {path.name} の転写結果が空でした。")
    elif not text:
        sys.exit(f"エラー: {path.name} からテキストを取得できませんでした。")
    return text


# ---------------------------------------------------------------------------
# 2. 機械チェック(決定論的な検証)
# ---------------------------------------------------------------------------
MONTHS = {
    "january": "1", "february": "2", "march": "3", "april": "4",
    "may": "5", "june": "6", "july": "7", "august": "8",
    "september": "9", "october": "10", "november": "11", "december": "12",
}


def extract_numbers(text: str) -> Counter:
    """比較用に数値を正規化して抽出する(全角→半角、月名→数字、カンマ除去)。"""
    t = text.translate(str.maketrans("0123456789", "0123456789"))
    for name, num in MONTHS.items():
        t = re.sub(name, f" {num} ", t, flags=re.IGNORECASE)
    t = re.sub(r"(?<=\d),(?=\d)", "", t)  # 300,000 → 300000
    return Counter(re.findall(r"\d+", t))


def extract_article_numbers(text: str) -> set[int]:
    """条見出し(行頭)のみを対象とする。文中の法令参照
    (例: 著作権法第27条)を誤検出しないよう行頭アンカーで判定する。"""
    if is_japanese(text):
        nums = re.findall(r"^第(\d+)条", text, flags=re.MULTILINE)
    else:
        nums = re.findall(r"^Articles?\s+(\d+)", text,
                          flags=re.MULTILINE | re.IGNORECASE)
    return {int(n) for n in nums}


def mechanical_checks(source: str, translation: str) -> list[str]:
    findings = []

    # (1) 条番号の欠落・過剰
    src_articles = extract_article_numbers(source)
    trans_articles = extract_article_numbers(translation)
    missing = sorted(src_articles - trans_articles)
    extra = sorted(trans_articles - src_articles)
    if missing:
        findings.append(f"[重大] 翻訳に存在しない条番号: {missing}(翻訳漏れの可能性)")
    if extra:
        findings.append(f"[中] 原文に存在しない条番号: {extra}")
    if not missing and not extra and src_articles:
        findings.append(f"[OK] 条番号は一致({len(src_articles)}箇条)")

    # (2) 数値の不一致(月名は数字に正規化して比較)
    src_nums = extract_numbers(source)
    trans_nums = extract_numbers(translation)
    missing_nums = sorted((src_nums - trans_nums).items(),
                          key=lambda x: -len(x[0]))[:15]
    if missing_nums:
        detail = ", ".join(f"{n}(×{c})" for n, c in missing_nums)
        findings.append(
            f"[要確認] 原文にあり翻訳に見つからない数値: {detail}\n"
            "        ※ 表記変換(漢数字・序数など)による誤検出の可能性もあります"
        )
    else:
        findings.append("[OK] 原文の数値はすべて翻訳に存在")

    # (3) 判読不能マーカーの残存(OCR版の出力チェック)
    markers = re.findall(r"\[illegible\]|\[判読不能\]", translation)
    if markers:
        findings.append(f"[中] 判読不能マーカーが {len(markers)} 箇所残っています")

    # (4) 分量バランス(極端に短い=大量の翻訳漏れの疑い)
    ratio = len(translation) / max(len(source), 1)
    findings.append(f"[情報] 文字数比(翻訳/原文): {ratio:.2f} "
                    f"(原文 {len(source):,} 字 → 翻訳 {len(translation):,} 字)")
    if (is_japanese(source) and not is_japanese(translation) and ratio < 1.2) or \
       (not is_japanese(source) and is_japanese(translation) and ratio < 0.25):
        findings.append("[要確認] 翻訳文が原文に対して短すぎます(翻訳漏れの疑い)")

    return findings


# ---------------------------------------------------------------------------
# 3. AIレビュー(Claude Opus)
# ---------------------------------------------------------------------------
def parse_json_response(raw: str) -> dict:
    cleaned = re.sub(r"^```(?:json)?|```$", "", raw.strip(),
                     flags=re.MULTILINE).strip()
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        # JSON部分だけを抽出して再試行
        m = re.search(r"\{.*\}", cleaned, flags=re.DOTALL)
        if m:
            try:
                return json.loads(m.group(0))
            except json.JSONDecodeError:
                pass
    return {"overall_assessment": "(JSON解析に失敗したため原文を表示)",
            "quality_score": None, "issues": [], "raw": raw}


def ai_review(source: str, translation: str, model: str,
              client: anthropic.Anthropic | None = None) -> dict:
    client = client or get_client()

    user_text = (REVIEW_PROMPT
                 .replace("{SOURCE}", source)
                 .replace("{TRANSLATION}", translation))

    last_error = None
    for attempt in range(1, MAX_RETRIES + 1):
        try:
            with client.messages.stream(
                model=model,
                max_tokens=MAX_TOKENS_PER_CALL,
                system=SYSTEM_PROMPT,
                messages=[{"role": "user", "content": user_text}],
            ) as stream:
                return parse_json_response(stream.get_final_text())
        except anthropic.AuthenticationError:
            sys.exit(
                "エラー: APIキーが無効です(401)。"
                ".env の ANTHROPIC_API_KEY を確認してください。"
            )
        except anthropic.APIError as e:
            last_error = e
            wait = 2 ** attempt
            print(f"  [警告] API エラー(試行 {attempt}/{MAX_RETRIES}): {e}")
            print(f"  {wait} 秒待機してリトライします...")
            time.sleep(wait)
    raise RuntimeError(f"AIレビューに失敗しました: {last_error}")


# ---------------------------------------------------------------------------
# 4. レポート生成
# ---------------------------------------------------------------------------
SEVERITY_ORDER = {"重大": 0, "中": 1, "軽微": 2}


def build_report(source_path: Path, trans_path: Path,
                 mech: list[str], ai: dict, model: str) -> str:
    lines = []
    lines.append("# 翻訳品質チェックレポート")
    lines.append("")
    lines.append(f"- 原文: `{source_path}`")
    lines.append(f"- 翻訳: `{trans_path}`")
    lines.append(f"- レビューモデル: `{model}`")
    lines.append("")

    lines.append("## 1. 機械チェック(条番号・数値・分量)")
    lines.append("")
    for f in mech:
        lines.append(f"- {f}")
    lines.append("")

    lines.append("## 2. AIレビュー(誤訳・翻訳漏れ・用語)")
    lines.append("")
    score = ai.get("quality_score")
    if score is not None:
        lines.append(f"**品質スコア: {score} / 100**")
        lines.append("")
    lines.append(f"**総評:** {ai.get('overall_assessment', '')}")
    lines.append("")

    issues = sorted(ai.get("issues", []),
                    key=lambda x: SEVERITY_ORDER.get(x.get("severity"), 9))
    if not issues:
        lines.append("指摘事項はありませんでした。")
    else:
        counts = Counter(i.get("severity", "?") for i in issues)
        lines.append(f"指摘件数: {len(issues)} 件"
                     f"(重大 {counts.get('重大', 0)} / 中 {counts.get('中', 0)} / "
                     f"軽微 {counts.get('軽微', 0)})")
        lines.append("")
        for n, issue in enumerate(issues, 1):
            lines.append(f"### {n}. [{issue.get('severity', '?')}] "
                         f"{issue.get('type', '?')} — {issue.get('location', '')}")
            lines.append("")
            if issue.get("source_excerpt"):
                lines.append(f"- 原文: {issue['source_excerpt']}")
            if issue.get("translation_excerpt"):
                lines.append(f"- 翻訳: {issue['translation_excerpt']}")
            lines.append(f"- 問題点: {issue.get('description', '')}")
            if issue.get("suggestion"):
                lines.append(f"- 修正案: {issue['suggestion']}")
            lines.append("")

    if "raw" in ai:
        lines.append("## (参考)AIの生レスポンス")
        lines.append("")
        lines.append("```")
        lines.append(ai["raw"])
        lines.append("```")

    return "\n".join(lines) + "\n"


# ---------------------------------------------------------------------------
# メイン処理
# ---------------------------------------------------------------------------
def clean_path(raw: str) -> str:
    s = raw.strip().strip('"').strip("'")
    s = s.replace("\\ ", " ")
    return os.path.expanduser(s)


def main() -> None:
    source_path = Path(clean_path(SOURCE_FILE))
    trans_path = Path(clean_path(TRANSLATED_FILE))
    for p, label in ((source_path, "SOURCE_FILE"), (trans_path, "TRANSLATED_FILE")):
        if not p.exists():
            sys.exit(f"エラー: ファイルが見つかりません: {p}\n"
                     f"  スクリプト先頭の {label} を確認してください。")

    output_dir = Path(clean_path(OUTPUT_DIR))
    output_dir.mkdir(parents=True, exist_ok=True)

    client_holder: dict = {"client": None}
    print(f"原文を読み込み中: {source_path}")
    source = load_document(source_path, client_holder)
    print(f"翻訳を読み込み中: {trans_path}")
    translation = load_document(trans_path, client_holder)

    total = len(source) + len(translation)
    if total > MAX_TOTAL_CHARS:
        print(f"[警告] 合計 {total:,} 文字と長大です。"
              "レビュー精度が落ちる場合は分割を検討してください。")

    print("\n--- 機械チェック ---")
    mech = mechanical_checks(source, translation)
    for f in mech:
        print(f"  {f}")

    print(f"\n--- AIレビュー({DEFAULT_MODEL})を実行中... ---")
    ai = ai_review(source, translation, DEFAULT_MODEL,
                   client=client_holder.get("client"))

    report = build_report(source_path, trans_path, mech, ai, DEFAULT_MODEL)
    report_path = output_dir / f"{trans_path.stem}_review.md"
    report_path.write_text(report, encoding="utf-8")

    issues = ai.get("issues", [])
    score = ai.get("quality_score")
    print(f"\nレビュー完了: 指摘 {len(issues)} 件"
          + (f" / 品質スコア {score}/100" if score is not None else ""))
    print(f"レポート: {report_path.resolve()}")


if __name__ == "__main__":
    main()