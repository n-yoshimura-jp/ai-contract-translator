#!/usr/bin/env python3
"""
日本語契約書 → 英語翻訳ツール
================================
PDF / Word (.docx) / テキスト (.txt) の契約書を読み込み、
Anthropic Claude API で英語に翻訳し、PDF / Word / テキストの3形式で出力します。

必要ライブラリ:
    pip install anthropic pdfplumber python-docx reportlab python-dotenv

APIキー設定(.env ファイル):
    スクリプトと同じディレクトリに .env ファイルを作成し、以下を記述:
        ANTHROPIC_API_KEY=sk-ant-...
    ※ .env は .gitignore に必ず追加してください(キー漏洩防止)

入出力パスの設定:
    スクリプト先頭の「★★★ 設定 ★★★」ブロックにある
    INPUT_FILE と OUTPUT_DIR を書き換えてください。

使い方:
    1. INPUT_FILE / OUTPUT_DIR を編集
    2. 実行: python translate_ja2en.py
"""

import os
import re
import sys
import time
from pathlib import Path

import anthropic
from dotenv import load_dotenv

# .env ファイルから環境変数を読み込む
# (スクリプトのあるディレクトリから親方向に .env を探索。
#  既存の環境変数が設定済みの場合はそちらを優先)
load_dotenv()

# ===========================================================================
# ★★★ 設定(ここを編集してください)★★★
# ===========================================================================
INPUT_FILE = "samples/sample_contract_ja.pdf"    # 入力ファイルのパス(.pdf / .docx / .txt)
OUTPUT_DIR = "output"        # 出力先ディレクトリのパス(無ければ自動作成)
# 例(絶対パスやユーザーホームも指定可能):
#   INPUT_FILE = "/Users/xxx/Documents/業務委託契約書.docx"
#   INPUT_FILE = "~/Desktop/契約書.txt"
#   OUTPUT_DIR = "~/Desktop/翻訳結果"
# ===========================================================================

# ---------------------------------------------------------------------------
# 内部設定
# ---------------------------------------------------------------------------
DEFAULT_MODEL = "claude-sonnet-5"
MAX_CHARS_PER_CHUNK = 8000   # 1回のAPI呼び出しで翻訳する最大文字数(日本語)
MAX_RETRIES = 3              # API失敗時のリトライ回数

SYSTEM_PROMPT = """You are a professional legal translator specializing in \
Japanese-to-English translation of contracts and legal documents.

Rules:
1. Translate the Japanese contract text into precise, natural legal English.
2. Preserve the document structure exactly: article numbers (第1条 → Article 1), \
clause numbering, paragraph breaks, and indentation.
3. Use standard legal terminology (e.g. 甲 → "Party A", 乙 → "Party B", \
本契約 → "this Agreement", 損害賠償 → "damages").
4. Do NOT summarize, omit, or add anything. Translate everything faithfully.
5. Output ONLY the translated text. No preamble, no commentary, no code fences.
6. If a segment starts or ends mid-sentence, translate it as-is without \
completing the sentence yourself."""


# ---------------------------------------------------------------------------
# 1. ファイル読み込み
# ---------------------------------------------------------------------------
def read_pdf(path: Path) -> str:
    import pdfplumber
    pages = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            pages.append(text)
    full_text = "\n\n".join(pages).strip()
    if not full_text:
        raise ValueError(
            "PDFからテキストを抽出できませんでした。"
            "スキャン画像のPDFの場合はOCR(例: pytesseract)が必要です。"
        )
    return full_text


def read_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts = []
    for para in doc.paragraphs:
        parts.append(para.text)
    # 表(テーブル)内のテキストも取得
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def read_txt(path: Path) -> str:
    # 日本語ファイルで一般的なエンコーディングを順に試す
    for enc in ("utf-8", "utf-8-sig", "cp932", "shift_jis", "euc-jp"):
        try:
            return path.read_text(encoding=enc).strip()
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise ValueError("テキストファイルの文字コードを判別できませんでした。")


def read_source(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return read_pdf(path)
    if suffix in (".docx", ".doc"):
        if suffix == ".doc":
            raise ValueError(".doc(旧形式)は未対応です。.docxに変換してください。")
        return read_docx(path)
    if suffix in (".txt", ".text", ".md"):
        return read_txt(path)
    raise ValueError(f"未対応のファイル形式です: {suffix}(PDF / .docx / .txt に対応)")


# ---------------------------------------------------------------------------
# 2. 日本語文字数カウント
# ---------------------------------------------------------------------------
JP_PATTERNS = {
    "漢字":       re.compile(r"[\u4E00-\u9FFF\u3400-\u4DBF\uF900-\uFAFF]"),
    "ひらがな":   re.compile(r"[\u3041-\u309F]"),
    "カタカナ":   re.compile(r"[\u30A0-\u30FF\u31F0-\u31FF\uFF66-\uFF9F]"),
    "全角記号等": re.compile(r"[\u3000-\u303F\uFF01-\uFF65]"),
}


def count_japanese_chars(text: str) -> dict[str, int]:
    """種類別の日本語文字数と全体の文字数を返す。"""
    counts = {label: len(p.findall(text)) for label, p in JP_PATTERNS.items()}
    counts["日本語文字 合計"] = sum(counts.values())
    counts["総文字数(空白・改行除く)"] = len(re.sub(r"\s", "", text))
    counts["総文字数(全体)"] = len(text)
    return counts


def print_char_stats(text: str) -> None:
    counts = count_japanese_chars(text)
    print("  --- 文字数カウント ---")
    for label, n in counts.items():
        print(f"    {label}: {n:,} 文字")
    print("  ----------------------")


# ---------------------------------------------------------------------------
# 3. 翻訳(Anthropic API)
# ---------------------------------------------------------------------------
def split_into_chunks(text: str, max_chars: int = MAX_CHARS_PER_CHUNK) -> list[str]:
    """段落境界を保ちながらテキストを分割する(長文契約書対応)。"""
    paragraphs = text.split("\n")
    chunks, current, current_len = [], [], 0
    for para in paragraphs:
        para_len = len(para) + 1
        if current and current_len + para_len > max_chars:
            chunks.append("\n".join(current))
            current, current_len = [], 0
        current.append(para)
        current_len += para_len
    if current:
        chunks.append("\n".join(current))
    return chunks


def translate_chunk(client: anthropic.Anthropic, chunk: str, model: str) -> str:
    last_error = None
    for attempt in range(1, MAX_RETRIES + 1):
        try:
            response = client.messages.create(
                model=model,
                max_tokens=8192,
                system=SYSTEM_PROMPT,
                messages=[{
                    "role": "user",
                    "content": (
                        "Translate the following Japanese contract text "
                        "into English:\n\n" + chunk
                    ),
                }],
            )
            return response.content[0].text.strip()
        except anthropic.AuthenticationError:
            sys.exit(
                "エラー: APIキーが無効です(401)。"
                ".env の ANTHROPIC_API_KEY を確認してください。"
            )
        except anthropic.APIError as e:
            last_error = e
            wait = 2 ** attempt
            print(f"  [警告] API エラー(試行 {attempt}/{MAX_RETRIES}): {e}")
            print(f"  {wait} 秒待機してリトライします...")
            time.sleep(wait)
    raise RuntimeError(f"翻訳に失敗しました: {last_error}")


def translate_text(text: str, model: str) -> str:
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        sys.exit(
            "エラー: ANTHROPIC_API_KEY が見つかりません。\n"
            "  スクリプトと同じディレクトリに .env ファイルを作成し、\n"
            "  次の1行を記述してください:\n"
            "    ANTHROPIC_API_KEY=sk-ant-..."
        )
    client = anthropic.Anthropic(api_key=api_key)
    chunks = split_into_chunks(text)
    print(f"翻訳を開始します({len(chunks)} チャンク / 約 {len(text):,} 文字)...")

    translated_parts = []
    for i, chunk in enumerate(chunks, 1):
        print(f"  チャンク {i}/{len(chunks)} を翻訳中...")
        translated_parts.append(translate_chunk(client, chunk, model))
    return "\n".join(translated_parts)


# ---------------------------------------------------------------------------
# 4. 出力(TXT / DOCX / PDF)
# ---------------------------------------------------------------------------
def write_txt(text: str, path: Path) -> None:
    path.write_text(text + "\n", encoding="utf-8")


def write_docx(text: str, path: Path) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    for line in text.split("\n"):
        line = line.rstrip()
        if not line:
            doc.add_paragraph("")
        elif line.lower().startswith("article ") or line.isupper():
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
        else:
            doc.add_paragraph(line)
    doc.save(str(path))


def write_pdf(text: str, path: Path) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    from xml.sax.saxutils import escape

    doc = SimpleDocTemplate(
        str(path), pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=20 * mm, bottomMargin=20 * mm,
    )
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "Body", parent=styles["Normal"],
        fontName="Times-Roman", fontSize=10.5, leading=15,
    )
    heading = ParagraphStyle(
        "ArticleHeading", parent=body,
        fontName="Times-Bold", spaceBefore=8, spaceAfter=4,
    )

    story = []
    for line in text.split("\n"):
        line = line.rstrip()
        if not line:
            story.append(Spacer(1, 6))
        else:
            style = heading if line.lower().startswith("article ") else body
            story.append(Paragraph(escape(line), style))
    doc.build(story)


# ---------------------------------------------------------------------------
# メイン処理
# ---------------------------------------------------------------------------
def clean_path(raw: str) -> str:
    """引用符や、Macのドラッグ&ドロップで付くエスケープを取り除く。"""
    s = raw.strip().strip('"').strip("'")
    s = s.replace("\\ ", " ")  # 例: /Users/xx/My\ Docs/契約書.pdf
    return os.path.expanduser(s)  # 「~/」をホームディレクトリに展開


def main() -> None:
    # --- コード先頭の設定値からパスを取得 ---
    input_path = Path(clean_path(INPUT_FILE))
    if not input_path.exists():
        sys.exit(
            f"エラー: ファイルが見つかりません: {input_path}\n"
            "  スクリプト先頭の INPUT_FILE を実際のファイルパスに書き換えてください。"
        )

    output_dir = Path(clean_path(OUTPUT_DIR))
    output_dir.mkdir(parents=True, exist_ok=True)

    # 1. 読み込み
    print(f"読み込み中: {input_path}")
    source_text = read_source(input_path)
    print("  抽出完了")
    print_char_stats(source_text)

    # 2. 翻訳
    translated = translate_text(source_text, DEFAULT_MODEL)

    # 3. 出力
    stem = input_path.stem + "_EN"
    outputs = {
        "テキスト": output_dir / f"{stem}.txt",
        "Word":     output_dir / f"{stem}.docx",
        "PDF":      output_dir / f"{stem}.pdf",
    }
    write_txt(translated, outputs["テキスト"])
    write_docx(translated, outputs["Word"])
    write_pdf(translated, outputs["PDF"])

    print("\n完了しました。出力ファイル:")
    for label, p in outputs.items():
        print(f"  [{label}] {p.resolve()}")


if __name__ == "__main__":
    main()