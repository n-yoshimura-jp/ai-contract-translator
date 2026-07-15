#!/usr/bin/env python3
"""
スキャン・画像入りPDF対応 契約書翻訳ツール【日本語 → 英語】
=====================================================================
PDFファイルをそのまま Anthropic API に送信し、Claudeのvision機能で
読み取り(テキスト+OCR)と英訳を1回のAPI呼び出しで行います。
スキャン画像PDF・画像入りPDF・通常のPDFのすべてに対応し、
結果を PDF / Word / テキストの3形式で出力します。

必要ライブラリ:
    pip install anthropic pdfplumber pypdf python-docx reportlab python-dotenv

APIキー設定(.env ファイル):
    スクリプトと同じディレクトリに .env ファイルを作成し、以下を記述:
        ANTHROPIC_API_KEY=sk-ant-...

使い方:
    1. 下の「★★★ 設定 ★★★」ブロックを編集
    2. 実行: python translate_ocr_ja2en.py

注意:
    ・入力はPDFのみ(Word・テキストは translate_ja2en.py を使用)
    ・APIの制約: 1リクエストあたり最大100ページ/32MB
    ・各ページは画像として処理されるため、テキスト送信より
      トークン消費(=料金)が多くなります
"""

import base64
import io
import os
import re
import sys
import time
from pathlib import Path

import anthropic
from dotenv import load_dotenv

load_dotenv()

# ===========================================================================
# ★★★ 設定(ここを編集してください)★★★
# ===========================================================================
INPUT_FILE = "samples/sample_contract_ja_scan.pdf"  # 入力PDFのパス(スキャン画像PDFもOK)
OUTPUT_DIR = "output"      # 出力先ディレクトリ(無ければ自動作成)
# ===========================================================================

# ---------------------------------------------------------------------------
# 内部設定
# ---------------------------------------------------------------------------
DEFAULT_MODEL = "claude-sonnet-5"
PAGES_PER_BATCH = 5          # 1回のAPI呼び出しで処理するページ数
MAX_TOKENS_PER_CALL = 32000  # 翻訳出力用のトークン上限
MAX_RETRIES = 3
MAX_BATCH_MB = 25            # 1バッチの上限サイズ(APIの32MB制限に余裕を持たせる)

OUTPUT_SUFFIX = "_EN"
HEADING_PATTERN = re.compile(r"^Article\s+\d+", re.IGNORECASE)

SYSTEM_PROMPT = (
    "You are a professional legal translator specializing in "
    "Japanese-to-English translation of contracts. You are also "
    "highly skilled at reading scanned documents accurately."
)

USER_PROMPT = (
    "The attached PDF contains a Japanese contract. It may be a "
    "scanned image without a text layer.\n\n"
    "Step 1: Read ALL Japanese text in the PDF accurately, including "
    "headings, numbered clauses, tables, dates, names, and addresses.\n"
    "Step 2: Translate everything into precise, natural legal English.\n\n"
    "Rules:\n"
    "- Preserve the document structure exactly "
    "(第1条 → Article 1, clause numbering, paragraph breaks).\n"
    "- Use standard legal terminology (甲 → Party A, 乙 → Party B, "
    "本契約 → this Agreement).\n"
    "- Do NOT summarize, omit, or add anything.\n"
    "- If a page starts or ends mid-sentence, translate it as-is.\n"
    "- If any part is illegible, write [illegible] at that position.\n"
    "- Output ONLY the English translation. No preamble, no commentary."
)


# ---------------------------------------------------------------------------
# 1. PDFの検査とページ分割
# ---------------------------------------------------------------------------
def count_pages(path: Path) -> int:
    import pdfplumber
    with pdfplumber.open(path) as pdf:
        return len(pdf.pages)


def has_text_layer(path: Path) -> bool:
    """テキスト層があるか簡易判定(情報表示用)。"""
    import pdfplumber
    with pdfplumber.open(path) as pdf:
        sample = pdf.pages[: min(3, len(pdf.pages))]
        text = "".join((p.extract_text() or "") for p in sample)
    return len(text.strip()) > 50


def split_pdf_into_batches(path: Path, pages_per_batch: int) -> list[bytes]:
    """PDFを指定ページ数ごとのバイト列に分割する。"""
    from pypdf import PdfReader, PdfWriter

    reader = PdfReader(str(path))
    total = len(reader.pages)
    batches = []
    for start in range(0, total, pages_per_batch):
        writer = PdfWriter()
        for page in reader.pages[start: start + pages_per_batch]:
            writer.add_page(page)
        buf = io.BytesIO()
        writer.write(buf)
        data = buf.getvalue()
        size_mb = len(data) / (1024 * 1024)
        if size_mb > MAX_BATCH_MB:
            sys.exit(
                f"エラー: ページ {start + 1}〜 のバッチが {size_mb:.1f}MB あり、"
                f"上限({MAX_BATCH_MB}MB)を超えます。\n"
                "  PAGES_PER_BATCH を小さくするか、PDFの解像度を下げてください。"
            )
        batches.append(data)
    return batches


# ---------------------------------------------------------------------------
# 2. 読み取り(テキスト+OCR)と翻訳を1回のAPI呼び出しで実行
# ---------------------------------------------------------------------------
def ocr_translate_batch(
    client: anthropic.Anthropic, pdf_bytes: bytes, model: str
) -> str:
    pdf_b64 = base64.standard_b64encode(pdf_bytes).decode("utf-8")
    last_error = None
    for attempt in range(1, MAX_RETRIES + 1):
        try:
            # max_tokensが大きい(処理が長時間になりうる)ためストリーミングで受信
            with client.messages.stream(
                model=model,
                max_tokens=MAX_TOKENS_PER_CALL,
                system=SYSTEM_PROMPT,
                messages=[{
                    "role": "user",
                    "content": [
                        {
                            "type": "document",
                            "source": {
                                "type": "base64",
                                "media_type": "application/pdf",
                                "data": pdf_b64,
                            },
                        },
                        {"type": "text", "text": USER_PROMPT},
                    ],
                }],
            ) as stream:
                return stream.get_final_text().strip()
        except anthropic.AuthenticationError:
            sys.exit(
                "エラー: APIキーが無効です(401)。"
                ".env の ANTHROPIC_API_KEY を確認してください。"
            )
        except anthropic.APIError as e:
            last_error = e
            wait = 2 ** attempt
            print(f"  [警告] API エラー(試行 {attempt}/{MAX_RETRIES}): {e}")
            print(f"  {wait} 秒待機してリトライします...")
            time.sleep(wait)
    raise RuntimeError(f"OCR・翻訳に失敗しました: {last_error}")


def ocr_translate_pdf(path: Path, model: str) -> str:
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        sys.exit(
            "エラー: ANTHROPIC_API_KEY が見つかりません。\n"
            "  .env ファイルに ANTHROPIC_API_KEY=sk-ant-... を記述してください。"
        )
    client = anthropic.Anthropic(api_key=api_key)

    total_pages = count_pages(path)
    layer = "あり(通常のPDF)" if has_text_layer(path) else "なし(スキャン画像PDF)"
    print(f"  ページ数: {total_pages} / テキスト層: {layer}")
    print("  ※ どちらの場合もvision機能で読み取ります")

    batches = split_pdf_into_batches(path, PAGES_PER_BATCH)
    print(f"読み取り+翻訳を開始します({len(batches)} バッチ、"
          f"最大 {PAGES_PER_BATCH} ページずつ)...")

    parts = []
    for i, pdf_bytes in enumerate(batches, 1):
        print(f"  バッチ {i}/{len(batches)} を処理中"
              f"({len(pdf_bytes) / 1024:.0f} KB)...")
        parts.append(ocr_translate_batch(client, pdf_bytes, model))
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# 3. 出力(TXT / DOCX / PDF)※ 英語出力用フォント
# ---------------------------------------------------------------------------
def write_txt(text: str, path: Path) -> None:
    path.write_text(text + "\n", encoding="utf-8")


def write_docx(text: str, path: Path) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    for line in text.split("\n"):
        line = line.rstrip()
        if not line:
            doc.add_paragraph("")
        elif HEADING_PATTERN.match(line):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
        else:
            doc.add_paragraph(line)
    doc.save(str(path))


def write_pdf(text: str, path: Path) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    from xml.sax.saxutils import escape

    doc = SimpleDocTemplate(
        str(path), pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=20 * mm, bottomMargin=20 * mm,
    )
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "Body", parent=styles["Normal"],
        fontName="Times-Roman", fontSize=10.5, leading=15,
    )
    heading = ParagraphStyle(
        "ArticleHeading", parent=body,
        fontName="Times-Bold", spaceBefore=8, spaceAfter=4,
    )

    story = []
    for line in text.split("\n"):
        line = line.rstrip()
        if not line:
            story.append(Spacer(1, 6))
        else:
            style = heading if HEADING_PATTERN.match(line) else body
            story.append(Paragraph(escape(line), style))
    doc.build(story)


# ---------------------------------------------------------------------------
# メイン処理
# ---------------------------------------------------------------------------
def clean_path(raw: str) -> str:
    s = raw.strip().strip('"').strip("'")
    s = s.replace("\\ ", " ")
    return os.path.expanduser(s)


def main() -> None:
    input_path = Path(clean_path(INPUT_FILE))
    if not input_path.exists():
        sys.exit(
            f"エラー: ファイルが見つかりません: {input_path}\n"
            "  スクリプト先頭の INPUT_FILE を実際のファイルパスに書き換えてください。"
        )
    if input_path.suffix.lower() != ".pdf":
        sys.exit(
            "エラー: このスクリプトはPDF専用です。\n"
            "  Word・テキストファイルは translate_ja2en.py を使用してください。"
        )

    output_dir = Path(clean_path(OUTPUT_DIR))
    output_dir.mkdir(parents=True, exist_ok=True)

    # 1. 検査 → 2. 読み取り+翻訳
    print(f"読み込み中: {input_path}")
    translated = ocr_translate_pdf(input_path, DEFAULT_MODEL)

    # 3. 出力
    stem = input_path.stem + OUTPUT_SUFFIX
    outputs = {
        "テキスト": output_dir / f"{stem}.txt",
        "Word":     output_dir / f"{stem}.docx",
        "PDF":      output_dir / f"{stem}.pdf",
    }
    write_txt(translated, outputs["テキスト"])
    write_docx(translated, outputs["Word"])
    write_pdf(translated, outputs["PDF"])

    print("\n完了しました。出力ファイル:")
    for label, p in outputs.items():
        print(f"  [{label}] {p.resolve()}")


if __name__ == "__main__":
    main()